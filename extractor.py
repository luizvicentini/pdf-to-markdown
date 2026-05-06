"""
Extração de PDF/TXT/DOCX em estrutura intermediária comum.

Saída: lista de blocos. Cada bloco é dict com:
  - kind: "heading" | "paragraph" | "list_item" | "table"
  - text: conteúdo textual
  - level: nível de heading (1, 2, 3) — apenas para headings
  - bold: bool — flag de negrito predominante
  - italic: bool — flag de itálico predominante
  - bbox: tupla (x0, y0, x1, y1) — apenas em PDF
  - page: índice da página — apenas em PDF
  - table_rows: list[list[str]] — apenas para tabelas
"""
import re
from collections import Counter
from pathlib import Path
from typing import List, Optional, Tuple

import fitz  # PyMuPDF


# Flags do PyMuPDF para spans de texto
FLAG_BOLD = 16
FLAG_ITALIC = 2

# Prefixos de lista
RE_BULLET = re.compile(r"^\s*[•\-\*•○▪]\s+")
RE_NUMBERED = re.compile(r"^\s*\d+[.)]\s+")


class ExtractionError(Exception):
    """Erro reportável ao usuário durante extração."""
    pass


# ----------------------------- PDF -----------------------------

def _is_bold(flags: int, font_name: str) -> bool:
    """Detecta negrito por flags ou nome da fonte."""
    if flags & FLAG_BOLD:
        return True
    name = (font_name or "").lower()
    return "bold" in name or "black" in name or "heavy" in name


def _is_italic(flags: int, font_name: str) -> bool:
    """Detecta itálico por flags ou nome da fonte."""
    if flags & FLAG_ITALIC:
        return True
    name = (font_name or "").lower()
    return "italic" in name or "oblique" in name


def _aggregate_block(block: dict) -> Tuple[str, float, bool, bool]:
    """
    Concatena spans de um bloco e devolve (texto, tamanho_predominante, bold, italic).
    Tamanho predominante = média ponderada por nº de chars.
    """
    parts = []
    sizes_weighted = []
    bold_count = 0
    italic_count = 0
    span_count = 0

    for line in block.get("lines", []):
        line_parts = []
        for span in line.get("spans", []):
            text = span.get("text", "")
            if not text:
                continue
            line_parts.append(text)
            size = span.get("size", 0.0)
            sizes_weighted.append((size, len(text)))
            flags = span.get("flags", 0)
            font = span.get("font", "")
            if _is_bold(flags, font):
                bold_count += 1
            if _is_italic(flags, font):
                italic_count += 1
            span_count += 1
        if line_parts:
            parts.append("".join(line_parts))

    text = "\n".join(parts)
    if sizes_weighted:
        total_chars = sum(c for _, c in sizes_weighted)
        if total_chars > 0:
            avg_size = sum(s * c for s, c in sizes_weighted) / total_chars
        else:
            avg_size = sizes_weighted[0][0]
    else:
        avg_size = 0.0

    bold = span_count > 0 and bold_count / span_count >= 0.6
    italic = span_count > 0 and italic_count / span_count >= 0.6
    return text, avg_size, bold, italic


def _detect_heading_level(size: float, body_size: float, size_thresholds: List[float]) -> Optional[int]:
    """
    Mapeia tamanho de fonte para nível de heading.
    size_thresholds: lista descendente de tamanhos únicos significativamente maiores que o body.
    Retorna 1, 2, 3 ou None.
    """
    if size <= body_size * 1.05:
        return None
    for idx, threshold in enumerate(size_thresholds[:3]):
        if size >= threshold * 0.98:
            return idx + 1
    return None


def _detect_list_item(text: str) -> bool:
    """Verifica se texto começa com marcador de lista."""
    return bool(RE_BULLET.match(text) or RE_NUMBERED.match(text))


def _strip_list_marker(text: str) -> Tuple[str, str]:
    """Remove marcador de lista. Devolve (marcador, conteúdo)."""
    m = RE_BULLET.match(text)
    if m:
        return "-", text[m.end():]
    m = RE_NUMBERED.match(text)
    if m:
        return "1.", text[m.end():]
    return "", text


def _identify_repeated_headers_footers(
    raw_blocks: List[dict], page_count: int
) -> set:
    """
    Marca textos que aparecem em y similar em 3+ páginas — provável cabeçalho/rodapé.
    Retorna conjunto de chaves (texto_normalizado, faixa_y) consideradas repetidas.
    """
    if page_count < 3:
        return set()

    counter: Counter = Counter()
    for blk in raw_blocks:
        txt = blk["text"].strip()
        if not txt or len(txt) > 200:
            continue
        # Faixa de y arredondada para 20px — agrupa posições similares
        y_band = round(blk["bbox"][1] / 20) * 20
        # Considera apenas topo (<100) ou fundo (último ~15% da página)
        page_height = blk.get("page_height", 0)
        if blk["bbox"][1] < 100 or blk["bbox"][3] > page_height * 0.85:
            counter[(txt, y_band)] += 1

    threshold = max(3, page_count // 3)
    return {k for k, v in counter.items() if v >= threshold}


def extract_pdf_structure(path: str) -> List[dict]:
    """
    Extrai estrutura de PDF nativo. Retorna lista de blocos com metadados.
    Levanta ExtractionError em casos previsíveis (senha, escaneado, corrompido).
    """
    try:
        doc = fitz.open(path)
    except Exception as exc:
        raise ExtractionError(f"Não foi possível abrir o arquivo: {exc}") from exc

    if doc.is_encrypted:
        # Tenta senha vazia (alguns PDFs aceitam)
        if not doc.authenticate(""):
            doc.close()
            raise ExtractionError(
                "PDF protegido por senha. Remova a proteção antes de converter."
            )

    raw_blocks: List[dict] = []
    table_blocks: List[dict] = []
    all_sizes: List[Tuple[float, int]] = []  # (size, char_count)
    total_chars = 0

    try:
        for page_idx, page in enumerate(doc):
            page_height = page.rect.height

            # Tabelas detectadas pela própria PyMuPDF
            try:
                tables = page.find_tables()
                table_bboxes = []
                for tbl in tables:
                    rows = tbl.extract()
                    if not rows:
                        continue
                    # Limpa células None
                    cleaned = [
                        [(cell or "").strip().replace("\n", " ") for cell in row]
                        for row in rows
                    ]
                    table_blocks.append({
                        "kind": "table",
                        "table_rows": cleaned,
                        "bbox": tuple(tbl.bbox),
                        "page": page_idx,
                    })
                    table_bboxes.append(tbl.bbox)
            except Exception:
                table_bboxes = []

            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:  # tipo 0 = texto
                    continue
                bbox = block.get("bbox", (0, 0, 0, 0))
                # Pula blocos dentro de tabelas
                if _bbox_overlaps_any(bbox, table_bboxes):
                    continue

                text, size, bold, italic = _aggregate_block(block)
                text = text.strip()
                if not text:
                    continue

                raw_blocks.append({
                    "text": text,
                    "size": size,
                    "bold": bold,
                    "italic": italic,
                    "bbox": bbox,
                    "page": page_idx,
                    "page_height": page_height,
                })
                all_sizes.append((size, len(text)))
                total_chars += len(text)
    finally:
        doc.close()

    if not raw_blocks and not table_blocks:
        raise ExtractionError(
            "Nenhum texto extraído. PDF pode ser escaneado (apenas imagem). "
            "Use OCR externo (ex: Tesseract, Adobe Acrobat) antes de converter."
        )

    # Tamanho de corpo predominante = moda ponderada por chars
    size_chars: dict = {}
    for size, chars in all_sizes:
        rounded = round(size, 1)
        size_chars[rounded] = size_chars.get(rounded, 0) + chars
    body_size = max(size_chars.items(), key=lambda kv: kv[1])[0] if size_chars else 12.0

    # Tamanhos significativos acima do body — candidatos a heading
    heading_sizes = sorted(
        {round(s, 1) for s, _ in all_sizes if s > body_size * 1.05},
        reverse=True,
    )

    # Cabeçalhos/rodapés repetitivos
    repeated = _identify_repeated_headers_footers(raw_blocks, len(all_sizes) and (raw_blocks[-1]["page"] + 1))

    # Classifica cada bloco
    classified: List[dict] = []
    for blk in raw_blocks:
        txt = blk["text"].strip()
        y_band = round(blk["bbox"][1] / 20) * 20
        if (txt, y_band) in repeated:
            continue  # descarta cabeçalho/rodapé recorrente

        level = _detect_heading_level(blk["size"], body_size, heading_sizes)
        if level is not None and len(txt) < 200:
            classified.append({
                "kind": "heading",
                "text": txt,
                "level": level,
                "bold": blk["bold"],
                "italic": blk["italic"],
                "bbox": blk["bbox"],
                "page": blk["page"],
            })
            continue

        if _detect_list_item(txt):
            marker, content = _strip_list_marker(txt)
            classified.append({
                "kind": "list_item",
                "text": content.strip(),
                "marker": marker,
                "bold": blk["bold"],
                "italic": blk["italic"],
                "bbox": blk["bbox"],
                "page": blk["page"],
            })
            continue

        classified.append({
            "kind": "paragraph",
            "text": txt,
            "bold": blk["bold"],
            "italic": blk["italic"],
            "bbox": blk["bbox"],
            "page": blk["page"],
        })

    # Mescla tabelas e blocos por página + posição vertical
    merged = classified + table_blocks
    merged.sort(key=lambda b: (b.get("page", 0), b.get("bbox", (0, 0, 0, 0))[1]))
    return merged


def _bbox_overlaps_any(bbox, others) -> bool:
    """Retorna True se bbox tem interseção significativa com qualquer bbox de others."""
    x0, y0, x1, y1 = bbox
    for ox0, oy0, ox1, oy1 in others:
        if x0 < ox1 and x1 > ox0 and y0 < oy1 and y1 > oy0:
            return True
    return False


# ----------------------------- TXT -----------------------------

def extract_txt_structure(path: str) -> List[dict]:
    """Lê arquivo .txt e divide em blocos (parágrafos por linhas em branco)."""
    p = Path(path)
    try:
        content = p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            content = p.read_text(encoding="latin-1")
        except OSError as exc:
            raise ExtractionError(f"Falha ao ler .txt: {exc}") from exc
    except OSError as exc:
        raise ExtractionError(f"Falha ao ler .txt: {exc}") from exc

    blocks: List[dict] = []
    for chunk in re.split(r"\n\s*\n", content):
        chunk = chunk.strip()
        if not chunk:
            continue
        if _detect_list_item(chunk.splitlines()[0]):
            for line in chunk.splitlines():
                line = line.strip()
                if not line:
                    continue
                marker, content_line = _strip_list_marker(line)
                blocks.append({
                    "kind": "list_item",
                    "text": content_line.strip(),
                    "marker": marker or "-",
                    "bold": False,
                    "italic": False,
                })
        else:
            blocks.append({
                "kind": "paragraph",
                "text": chunk,
                "bold": False,
                "italic": False,
            })
    if not blocks:
        raise ExtractionError("Arquivo .txt está vazio.")
    return blocks


# ----------------------------- DOCX -----------------------------

def extract_docx_structure(path: str) -> List[dict]:
    """Extrai estrutura de .docx via python-docx. Mapeia estilos de heading e tabelas."""
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ExtractionError("Pacote python-docx não instalado.") from exc

    try:
        doc = docx.Document(path)
    except Exception as exc:
        raise ExtractionError(f"Falha ao abrir .docx: {exc}") from exc

    blocks: List[dict] = []

    # Itera parágrafos do corpo respeitando ordem
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()
        # Detecta heading pelo estilo (Heading 1, 2, 3...)
        m = re.match(r"heading\s*(\d+)", style_name)
        if m:
            level = min(int(m.group(1)), 3)
            blocks.append({
                "kind": "heading",
                "text": text,
                "level": level,
                "bold": True,
                "italic": False,
            })
            continue

        if _detect_list_item(text) or "list" in style_name:
            marker, content = _strip_list_marker(text)
            blocks.append({
                "kind": "list_item",
                "text": content.strip() if marker else text,
                "marker": marker or "-",
                "bold": False,
                "italic": False,
            })
            continue

        # Negrito/itálico predominante: checa runs
        bold_chars = 0
        italic_chars = 0
        total_chars = 0
        for run in para.runs:
            n = len(run.text)
            total_chars += n
            if run.bold:
                bold_chars += n
            if run.italic:
                italic_chars += n
        bold = total_chars > 0 and bold_chars / total_chars >= 0.6
        italic = total_chars > 0 and italic_chars / total_chars >= 0.6

        blocks.append({
            "kind": "paragraph",
            "text": text,
            "bold": bold,
            "italic": italic,
        })

    # Tabelas no final (python-docx não fornece ordem entre parágrafos e tabelas trivialmente)
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
        if rows:
            blocks.append({
                "kind": "table",
                "table_rows": rows,
            })

    if not blocks:
        raise ExtractionError("Documento .docx está vazio.")
    return blocks


# ----------------------------- Dispatcher -----------------------------

def extract_any(path: str) -> List[dict]:
    """Despacha para extrator correto pela extensão do arquivo."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return extract_pdf_structure(path)
    if ext == ".txt":
        return extract_txt_structure(path)
    if ext == ".docx":
        return extract_docx_structure(path)
    raise ExtractionError(f"Formato não suportado: {ext}")
